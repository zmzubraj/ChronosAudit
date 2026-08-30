from pathlib import Path
import json, math
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parent
ART=ROOT/'artifact'
OUT=ROOT/'paper'/'ChronosAudit_PreIncident_TopJournal_Stage2_ExternalEvidence_Revised.docx'
OUT.parent.mkdir(exist_ok=True)
audit=json.loads((ART/'reports'/'stage2a_2e_audit.json').read_text())
split=json.loads((ART/'reports'/'split_baseline_audit.json').read_text())
inc=pd.read_csv(ART/'raw'/'incident_evidence_enriched.csv')
seed=pd.read_csv(ART/'raw'/'scone_bench.csv')


def shade(cell, fill='EAEAEA'):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); h=OxmlElement('w:tblHeader'); h.set(qn('w:val'),'true'); trPr.append(h)

def cant_split(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:cantSplit'); trPr.append(e)

def body(doc,text,indent=True):
    p=doc.add_paragraph(text,style='Body Text')
    if indent: p.paragraph_format.first_line_indent=Inches(.22)
    return p

def bullet(doc,text):
    p=doc.add_paragraph(text,style='List Bullet'); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.14); return p

def caption(doc,text):
    p=doc.add_paragraph(style='Caption'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(text); return p

def table(doc,headers,rows,font=8.1,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    repeat_header(t.rows[0]); cant_split(t.rows[0])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); shade(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(font)
    for row in rows:
        cells=t.add_row().cells; cant_split(t.rows[-1])
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.size=Pt(font)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

def add_figure(doc,path,cap,alt,width=6.6):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run(); shape=run.add_picture(str(path),width=Inches(width))
    try:
        docPr=shape._inline.docPr; docPr.set('descr',alt); docPr.set('title',cap)
    except Exception: pass
    caption(doc,cap)

def equation(doc,text,num):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.name='Cambria Math'; r.font.size=Pt(10.5)
    p.add_run(f'    ({num})')


doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.78); sec.right_margin=Inches(.78)
sec.header_distance=Inches(.3); sec.footer_distance=Inches(.35)
styles=doc.styles
styles['Normal'].font.name='Times New Roman'; styles['Normal'].font.size=Pt(10.5)
styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.08
styles['Body Text'].font.name='Times New Roman'; styles['Body Text'].font.size=Pt(10.5); styles['Body Text'].paragraph_format.space_after=Pt(5); styles['Body Text'].paragraph_format.line_spacing=1.08
for name,size in [('Title',18),('Heading 1',13.5),('Heading 2',11.5),('Heading 3',10.7)]:
    st=styles[name]; st.font.name='Arial'; st.font.size=Pt(size); st.font.bold=True; st.paragraph_format.space_before=Pt(9); st.paragraph_format.space_after=Pt(4)
styles['Caption'].font.name='Times New Roman'; styles['Caption'].font.size=Pt(8.5); styles['Caption'].font.italic=True
styles['List Bullet'].font.name='Times New Roman'; styles['List Bullet'].font.size=Pt(10.5)
# header/footer
h=sec.header.paragraphs[0]; h.text='ChronosAudit | Anonymous manuscript | revised Stage-2 submission'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in h.runs: r.font.name='Arial'; r.font.size=Pt(8)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=f.add_run('Fail-closed evidence claims; no unexecuted live gate is reported as complete'); rr.font.name='Arial'; rr.font.size=Pt(8)

p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('ChronosAudit: Information-Admissible Cohorts and Contamination-Aware Evaluation for Pre-Incident Smart-Contract Exploit Detection')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Anonymous manuscript').bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Empirical software engineering / blockchain security research article'); r.italic=True

# Abstract
doc.add_heading('Abstract',1)
body(doc,(
'Pre-incident smart-contract exploit detection is easy to overstate retrospectively because the evaluator can observe artifacts that a defender did not possess before an incident: post-incident source verification, exploit transactions, upgraded proxy implementations, public root-cause labels, duplicated deployments, and related protocol or mechanism families. We present ChronosAudit, a real-world evidence and evaluation infrastructure that treats every detector input as a time-indexed admissibility claim. The system separates detector-visible evidence from evaluator-only outcomes; fixes an outcome-independent primary prediction landmark 24 h after deployment with preregistered 1 h, 7 d, and 30 d sensitivity landmarks; reconstructs canonical historical bytecode and proxy identity with block-hash-pinned multi-provider RPC reads; records first-defensible source availability; creates blinded dual-review packets and immutable third-adjudication records; constructs a deployment denominator and cutoff-safe matched controls; preserves right censoring; and certifies R0-R5 partitions over exact identity, bytecode clone, proxy implementation, protocol family, and exploit-mechanism family. In the executed public-evidence layer, all 417 historical exploit tasks are linked to content-hashed incident records, yielding 408 unique chain-address identities and nine repeated identity groups. Exact-identity leakage occurs in 100% of 1,000 independent-random, balanced shuffled K-fold, and chain-stratified K-fold simulations, while identity-group K-fold has zero crossings. The closed-form expectation for nine duplicate pairs under independent five-fold assignment is 7.2 crossing groups and a 99.9999488% probability of at least one crossing, demonstrating that the simulation is a validation check rather than the substantive novelty claim. The revised implementation passes 34 automated tests and adds cutoff-safe matching, blinded review, third-adjudicator ingestion, multiple trace backends, retry/backoff, statistical utilities, and deterministic split auditing. ChronosAudit deliberately releases no detector-effectiveness cohort until live dual-provider historical evidence, independently reviewed labels, a >=20,000-contract deployment denominator, >=4,170 matched controls, longitudinal outcomes, and independent external regeneration pass frozen gates. The contribution is therefore a production-oriented measurement system for making future pre-incident detector-effectiveness claims falsifiable, auditable, and resistant to information leakage.'),False)
p=doc.add_paragraph(); p.add_run('Keywords: ').bold=True; p.add_run('smart-contract security; pre-incident detection; temporal leakage; benchmark contamination; historical blockchain state; source provenance; censoring; reproducibility')

# 1
doc.add_heading('1. Introduction',1)
for txt in [
'Smart contracts increasingly mediate financial and governance workflows whose failures can produce immediate, irreversible losses. A detector used in practice must therefore answer a temporal question: could it have produced a useful warning before an exploit became known, using only information actually available at that time? Retrospective evaluation often answers a weaker question. Once an incident occurs, source code may be verified or republished, exploit transactions reveal state transitions, proxy implementations may be upgraded, public post-mortems localize the root cause, and benchmark curators may unintentionally place related contracts in both training and test data. Performance obtained under those conditions is not automatically evidence of pre-incident capability.',
'Existing work already establishes that realism and data hygiene materially change vulnerability-detection results. PrimeVul shows that aggressive deduplication and chronological splitting can collapse code-language-model performance relative to conventional vulnerability datasets [1]. DAppSCAN demonstrates the difficulty of evaluating smart-contract analyzers on large collections derived from real audit reports [2], and SmartBugs exposes substantial disagreement among automated analyzers [3]. DIVE explicitly reports repeated opcode skeletons and warns about train/test leakage in smart-contract data [4]. Recent agent benchmarks move closer to operational reality: EVMbench evaluates detection, exploitation and patching [9], ReEVMbench introduces post-model-release incidents to reduce model-training contamination [10], SCONE evaluates exploitation on historical state [11], and CyberChainBench anchors hundreds of real incidents to historical execution environments [12]. These advances leave a complementary measurement problem unresolved: historical executability is not the same as historical information admissibility.',
'ChronosAudit treats benchmark construction as a chain of evidence claims. The atomic unit is a specific chain-address deployment observed at a frozen prediction cutoff. Every detector-visible artifact must have a defensible availability time, provenance record, integrity hash and lineage status. Evaluator-only information such as exploit transactions and post-incident labels is permitted for outcome adjudication but cannot leak into detector inputs. If a required field cannot be established, the case remains in the audit ledger but is excluded from detector-effectiveness claims. This fail-closed design prevents a benchmark from silently becoming easier as more information appears after the incident.',
'This article reports the complete design and executable implementation of the Stage-2 evidence/cohort layer required by the broader ChronosAudit program. It does not manufacture detector performance in the absence of a scientifically qualified cohort. Instead, it closes the implementation defects identified in prior internal review, freezes the statistical and temporal policies that later stages depend on, and reports exactly which claims are already supported by executed public evidence and which still require live independent evidence. This distinction is essential for top-journal empirical work: a reproducible negative release decision is preferable to a populated benchmark whose information boundary cannot be defended.'
]: body(doc,txt)
body(doc,'The paper makes six contributions:',False)
for txt in [
'A contract-at-cutoff information-admissibility model with an outcome-independent 24-hour primary landmark and preregistered 1-hour, 7-day and 30-day sensitivity landmarks.',
'A real-world historical reconstruction implementation using canonical block-hash pinning, independent-provider consensus, common proxy resolution, multiple trace backends, source-history evidence, response hashing, and fail-closed provider-family independence checks.',
'A blinded two-reviewer plus third-adjudicator workflow with raw agreement, Cohen kappa, Gwet AC1, bootstrap confidence intervals, immutable final-decision hashes, and explicit separation of machine candidates from first-pass human judgments.',
'A cutoff-safe denominator/control design in which controls must already exist at the positive case cutoff; no future activity or outcome is allowed into matching. The design preserves unresolved/right-censored outcomes and defines IPCW and partial-identification analyses.',
'A contamination audit that compares independent random, balanced shuffled, chain-stratified and identity-group K-fold splitting, and generalizes the release protocol to R0-R5 family separation.',
'An executable, versioned artifact that currently links all 417 seed incidents to hashed public chronology while refusing to release detector-effectiveness cases until historical, reviewer, denominator, control, longitudinal and external-reproduction gates are actually satisfied.'
]: bullet(doc,txt)
add_figure(doc,ROOT/'figures/figure_1_stage2_pipeline.png','Figure 1. ChronosAudit Stage-2 evidence/cohort qualification pipeline.','Flow diagram showing Stage 2A through 2E and a fail-closed release gate.')

# 2
doc.add_heading('2. Research objective, estimands and completion gate',1)
body(doc,'The target claim is deliberately narrower than “zero-day prediction.” ChronosAudit asks whether a detector, supplied only with artifacts admissible at a frozen pre-incident cutoff, can discriminate later exploited contracts from contemporaneous controls and retain capability as increasingly strict contamination constraints are imposed. The primary scientific estimand in later detector stages is the change in precision-recall performance and analyst workload between conventional and contamination-controlled partitions, with uncertainty clustered at the strongest available dependency family. Stage 2 is complete only when the cohort itself is scientifically defensible; detector scores are not computed from cases that fail mandatory evidence gates.')
table(doc,['RQ','Question','Stage-2 completion criterion'],[
('RQ1','Can deployment, code, source and identity evidence be reconstructed as they existed at a frozen pre-incident cutoff?','Two independent provider families agree on canonical historical state; source availability is positively evidenced or marked unknown.'),
('RQ2','Can protocol and exploit-mechanism families be labeled without machine anchoring or single-reviewer authority?','Two blinded reviewers complete all cases; disagreements receive a third adjudicator; agreement statistics and confidence intervals are reported.'),
('RQ3','Can positives be embedded in a real deployment risk set without future-information leakage?','>=20,000 real deployments are collected; >=10 cutoff-safe matched controls per positive are frozen; unresolved outcomes remain censored.'),
('RQ4','Can train/test dependency be eliminated successively from exact identity through held-out mechanism family?','R0-R5 contamination graph is complete; required leakage is zero; >=120 independent R5 blocks remain.'),
('RQ5','Can another group regenerate the released cohort and partitions from frozen public inputs and provider observations?','Independent re-execution reproduces cohort membership and partition hashes within preregistered tolerances.')
])

# 3 related work
doc.add_heading('3. Related work and scholarly positioning',1)
doc.add_heading('3.1 Static, symbolic and learning-based smart-contract analysis',2)
body(doc,'Smart-contract security research includes static analysis, symbolic execution, fuzzing and learned representations. Slither provides a widely used static-analysis framework and intermediate representation for Solidity [5]. SmartBugs integrates multiple analyzers and, in a large empirical study, showed both limited shared detection and substantial false-positive concerns [3]. GPTScan combines LLM-guided candidate reasoning with program-analysis confirmation for logic vulnerabilities [6]. MANDO learns heterogeneous control-flow and call-graph representations at contract and line level [7], while MANDO-LLM combines heterogeneous graph transformers with LLM-derived features [8]. These systems address how to detect vulnerabilities; ChronosAudit addresses whether an evaluation gives those systems only information a defender could legitimately have possessed before the incident.')
doc.add_heading('3.2 Dataset quality, deduplication and temporal contamination',2)
body(doc,'Dataset contamination is not a smart-contract-specific concern. PrimeVul shows that label quality, duplication and unrealistic splits can substantially inflate vulnerability-detection estimates, with large performance degradation after stronger deduplication and chronological evaluation [1]. DAppSCAN builds source- and bytecode-level smart-contract datasets from 1,199 audit reports and 9,154 weaknesses, emphasizing realistic project data [2]. DIVE contributes a multi-label smart-contract dataset and explicitly identifies repeated opcode skeletons as a leakage risk [4]. ChronosAudit extends this line of work from duplicate functions or opcode structure to a broader dependency graph: chain-address identity, normalized bytecode, proxy implementation lineage, protocol/entity family and exploit-mechanism family.')
doc.add_heading('3.3 Historical exploit and agent benchmarks',2)
body(doc,'EVMbench evaluates AI agents on detection, exploitation and patching across curated vulnerabilities [9]. ReEVMbench demonstrates that model/scaffold rankings can change when evaluation uses 22 incidents that postdate model releases, highlighting training-data contamination and scaffold choice as empirical threats [10]. SCONE evaluates smart-contract exploitation in historical simulation and includes post-knowledge-cutoff tasks [11]. CyberChainBench expands real-world agent evaluation to 541 incidents across nine EVM chains with historical blockchain state and structured ground truth [12]. Chaliasos et al. evaluate five security tools against 127 high-impact real attacks and find that only a small minority were preventable by those tools [29]. Bastet contributes 849 fully expert-annotated DeFi findings produced by a two-annotator consensus workflow, providing a useful external reference for human-label reliability [30]. These benchmarks improve task realism and label quality; ChronosAudit targets a distinct validity dimension. A contract may be executable at an old block while its source, proxy implementation annotation, protocol relationship or exploit-family label became known only later. Historical state is therefore necessary but not sufficient for a pre-incident claim.')
doc.add_heading('3.4 Gap summary',2)
table(doc,['Capability','Representative prior work','ChronosAudit addition'],[
('Real-world smart-contract weaknesses','DAppSCAN, SmartBugs, DIVE [2-4]','Time-index every detector artifact and preserve unknown availability rather than substituting current artifacts.'),
('Graph/LLM detection','MANDO, MANDO-LLM, GPTScan [6-8]','Detector-agnostic admissibility and family-controlled evaluation.'),
('Historical execution','SCONE, CyberChainBench [11-12]','Block-hash-pinned multi-provider identity plus source-at-cutoff and proxy-lineage admissibility.'),
('Agent capability','EVMbench, ReEVMbench [9-10]','Explicit model-independent contamination ladder and release gate.'),
('General vulnerability leakage','PrimeVul [1]','Extend deduplication from code/time to deployment, proxy, protocol and mechanism families.'),
('Dataset documentation','Datasheets [24]','Hash-linked evidence registry, reason-coded exclusions, external-evidence gate ledger and deterministic regeneration.')
])

# 4 threat model
doc.add_heading('4. Threat model and information-admissibility model',1)
body(doc,'The adversary in the validity analysis is not only an exploit attacker; it is also accidental benchmark leakage. Leakage can occur when a detector receives information that was unavailable at the cutoff, when related contracts straddle train/test partitions, or when an apparently negative control is labeled safe despite insufficient follow-up. ChronosAudit therefore distinguishes two information planes. The detector plane contains only artifacts with positive pre-cutoff admissibility evidence. The evaluator plane may contain later incident reports, exploit transactions and adjudication notes, but those fields are physically and logically excluded from detector adapters.')
add_figure(doc,ROOT/'figures/figure_2_information_planes.png','Figure 2. Detector-visible and evaluator-only information planes.','Diagram separating pre-cutoff detector inputs from post-incident evaluator-only evidence.')
equation(doc,'t_deploy < t_cut < t_incident',1)
body(doc,'The primary prediction cutoff is frozen from deployment time alone: the first canonical observation at or after deployment + 24 hours. Incident time is used only after the cutoff is fixed, to determine whether the case has at least one hour of lead time. Sensitivity analyses repeat the cohort construction at deployment + 1 hour, +7 days and +30 days. Incidents that occur before a landmark are not moved to a convenient pre-incident time; they are ineligible for that landmark and may enter an earlier preregistered sensitivity stratum. This rule prevents retrospective knowledge of the exploit date from moving the cutoff.')
table(doc,['Landmark','Definition','Purpose'],[
('Primary','deployment + 24 h; require >=1 h lead before incident','Outcome-independent main analysis.'),
('Early','deployment + 1 h','Captures rapidly exploited deployments.'),
('Medium','deployment + 7 d','Tests stability after short operational exposure.'),
('Long','deployment + 30 d','Tests mature-contract evidence and survivor effects.')
])

# 5 cohort
doc.add_heading('5. Real-world positive cohort and incident evidence',1)
body(doc,'The frozen seed contains 417 SCONE tasks sourced from DeFiHackLabs: 226 on BNB Smart Chain, 181 on Ethereum mainnet, 9 on Base and 1 on Arbitrum. The enrichment pipeline maps every task to a content-hashed public incident record, preserving the source snapshot hash, incident-record hash, public incident name/date, raw public mechanism description and any transaction-hash hints. Coverage of the public chronology layer is 417/417. This does not mean 417/417 incidents are independently adjudicated ground truth; it means every seed row has a reproducible public incident linkage that can be independently checked.')
add_figure(doc,ROOT/'figures/figure_6_chain_composition.png','Figure 3. Chain composition of the 417-case seed.','Bar chart showing BNB Smart Chain, Ethereum, Base and Arbitrum composition.')
body(doc,'Cohort selection is reported as a flow rather than treating the seed as representative by assumption. The seed is conditioned on inclusion in SCONE/DeFiHackLabs and therefore may overrepresent incidents with reproducible exploit artifacts, public post-mortems or contracts that are easy to fork. ChronosAudit does not use the seed to estimate ecosystem exploit prevalence. Prevalence-sensitive metrics use the separate deployment denominator, and detector conclusions are stratified by chain, protocol family, mechanism family, proxy status, source availability and calendar period. This separates “known exploited case collection” from “real deployment risk set.”')

# 6 historical/source
doc.add_heading('6. Historical on-chain reconstruction and source-at-cutoff evidence',1)
body(doc,'Historical identity is reconstructed at the canonical cutoff block using at least two independently operated archive-provider families. Providers first agree on the block hash; subsequent state reads use an EIP-1898 block-hash selector so a same-height reorganization or provider inconsistency cannot silently change the observation [18]. The implementation hashes raw provider responses and normalizes runtime bytecode before computing identity fingerprints. EIP-1967 implementation, admin and beacon slots are read at the same block; beacon implementation() is resolved historically; EIP-1167 minimal proxies are recognized from runtime bytecode [19-20]. Diamond and custom/metamorphic patterns remain fail-closed unless a dedicated resolver can prove their historical implementation set.')
body(doc,'Provider independence is defined by operator family rather than URL count. Current documentation indicates archive support for Ethereum, BNB Smart Chain, Base and Arbitrum from QuickNode, and Chainstack provides an independent archive-provider family [15-16]. The collector records provider family, block hash, trace backend and response hashes. A chain does not pass because two endpoints respond; it passes only after capability probes demonstrate historical block/state access and at least two verified provider families agree on the required evidence.')
body(doc,'Source admissibility is asymmetric. Sourcify v2 exposes verified-contract metadata and a verification timestamp that can provide positive evidence that a source artifact existed no later than the cutoff [13]. Etherscan V2 is used as an independent current source/deployment cross-check and provides creator, creation transaction, block number, timestamp and creation bytecode, but its current source endpoint does not by itself establish the first public-availability time [14]. Therefore source_verified_at_cutoff is True only on positive timestamped evidence; a later or missing Sourcify timestamp is not interpreted as proof that the source was unavailable everywhere. Bytecode-only detectors can enter a bytecode-admissible cohort even when source availability is unknown, preventing source-verification bias from being imposed on all detector families.')
table(doc,['Cohort view','Required detector-visible evidence','Use'],[
('Bytecode-admissible','canonical runtime bytecode at cutoff; deployment/lineage evidence','Bytecode/static EVM models.'),
('Source-admissible','bytecode-admissible + positive source availability by cutoff','Solidity/source models.'),
('State-admissible','bytecode-admissible + required historical state observations','State-aware symbolic/agent systems.'),
('Source+state','all above','Richest detector plane; never used as universal denominator.')
])

# 7 review
doc.add_heading('7. Independent protocol/mechanism adjudication',1)
body(doc,'Public incident labels are useful retrieval cues but are not treated as adjudicated root causes. The first reviewer pass is now blinded to machine-generated protocol and mechanism candidates. Each reviewer receives only case identity, public incident evidence, hashes and references. Reviewer A and Reviewer B independently assign protocol family, primary root cause, confidence and evidence references. Only after both label sets are frozen are automated candidates revealed for discrepancy analysis. Cases with protocol or mechanism disagreement are routed to a third adjudicator, whose final labels, rationale and evidence references are ingested and hashed into an immutable final-decision record.')
body(doc,'Agreement reporting is deliberately redundant. The implementation computes raw agreement, Cohen kappa [21], Gwet AC1 because kappa can be unstable under skewed prevalence [22], and deterministic bootstrap 95% confidence intervals. Per-family confusion matrices and low-confidence rates are retained for the submission artifact. The preregistered minimum agreement threshold is kappa >=0.80 and AC1 >=0.80 for both protocol and primary-mechanism labels; if either measure or its uncertainty indicates instability, the family taxonomy is revised before release rather than relaxing the threshold post hoc.')

# 8 controls/censor
doc.add_heading('8. Deployment denominator, matched controls and censor-aware outcomes',1)
body(doc,'The denominator is a real deployment stream, not a synthetic negative set. Top-level contract creations are recovered from creation transactions and receipts. Internal CREATE/CREATE2 events are recovered through either Parity-style trace_block or Geth-style debug_traceBlockByNumber with callTracer; provider capability probing selects the available backend and records backend identity. A block slice is complete only when two verified provider families agree on the canonical block and creation record set. This addresses the portability defect in the earlier collector, which assumed trace_block everywhere.')
body(doc,'The primary denominator target is at least 20,000 deployments across the represented chains and calendar windows. This threshold is a precision floor, not a claim that 20,000 is universally optimal: at a 1% alert rate, 20,000 independent observations give approximately +/-0.14 percentage-point 95% Wilson precision before clustering. For 417 positives, the frozen 10:1 matched-control target yields 4,170 controls; at a 1% false-positive rate this corresponds to approximately +/-0.31 percentage-point 95% Wilson precision before design effects. Clustered confidence intervals and sensitivity analyses account for the fact that deployments within protocols and implementation families are not independent.')
body(doc,'Control selection is risk-set and cutoff safe. Each positive record must carry prediction_cutoff_time. A candidate control must be deployed on or before that cutoff, lie within the preregistered deployment-time caliper, satisfy code-size bounds, and—when enabled—match proxy and source-availability strata. The matcher never uses post-cutoff transaction activity, exploit outcome or later source publication. Deterministic SHA-256 ranking resolves ties. Controls that are not observed to be exploited during available follow-up are not called “safe”; their outcomes are unresolved/right-censored unless the target property and follow-up horizon justify a stronger label.')
equation(doc,'w_i = delta_i / G_hat(t_i | X_i)',2)
body(doc,'For censoring-sensitive analyses, inverse-probability-of-censoring weighting (IPCW) uses the estimated censoring survival G_hat(t|X). The confirmatory analysis also reports worst/best-case bounds for unresolved outcomes, landmark-specific follow-up, and complete-case sensitivity. This makes the detector conclusion robust to the fact that “not yet exploited” is not equivalent to “non-vulnerable.”')

# 9 R0-R5
doc.add_heading('9. R0-R5 contamination ladder and split certification',1)
body(doc,'ChronosAudit separates increasingly demanding generalization claims rather than collapsing them into one train/test split. R0 intentionally approximates conventional evaluation; subsequent levels remove stronger dependencies. A detector can therefore be described as recognizing duplicated implementations, transferring within protocols, generalizing across implementation families, or detecting genuinely held-out mechanism families, depending on where performance survives.')
table(doc,['Level','Constraint added','Interpretation'],[
('R0','Conventional benchmark-native or row-level split','Upper-bound conventional estimate; contamination may remain.'),
('R1','Exact chain-address and normalized bytecode groups cannot cross','Removes direct identity/clone recognition.'),
('R2','Prediction time strictly precedes test incidents; training artifacts obey time admissibility','Removes future temporal knowledge.'),
('R3','Protocol/entity families cannot cross','Tests cross-protocol transfer.'),
('R4','Proxy/implementation lineage and code-clone families cannot cross','Tests independence from reused implementation families.'),
('R5','Exploit-mechanism families cannot cross','Tests held-out-mechanism generalization.')
])
body(doc,f"The current seed contains {split['repeated_identity_groups']} repeated exact-identity groups, all pairs. We therefore report both simulation and the closed-form null expectation. Under independent five-fold assignment, the expected number of crossing duplicate pairs is {split['pair_only_theory']['expected_crossing_groups_independent_random']:.1f}, and the probability of at least one crossing is {100*split['pair_only_theory']['probability_at_least_one_crossing_independent_random']:.7f}%. Across 1,000 simulations, leakage occurred in 100% of independent-random, balanced shuffled K-fold, and chain-stratified balanced K-fold assignments; deterministic identity-group K-fold produced zero crossings. Because this result is mathematically expected for the observed pair structure, it is treated as a split-audit validation result rather than the paper's main novelty claim. The substantive future detector experiment is the R0-to-R5 capability-survival curve after bytecode, proxy, protocol and mechanism labels are fully adjudicated.")
add_figure(doc,ROOT/'figures/figure_8_split_strategy_leakage.png','Figure 4. Exact-identity leakage across row-level and identity-grouped split strategies.','Bar chart showing 100 percent leakage for three row-level strategies and zero for identity-group K-fold.')

# 10 stats
doc.add_heading('10. Statistical methodology and confirmatory analysis',1)
body(doc,'The primary detector-effectiveness analysis is preregistered before running detector systems on the release cohort. Precision-recall measures are primary because exploitation is rare and class imbalance makes ROC summaries potentially misleading [23]. We report precision, recall, PR-AUC, Brier score/calibration, abstention rate, alerts per 1,000 deployments, analyst minutes per true positive, and lead time. Detector thresholds are frozen on training/validation data and are not retuned on R5 or prospective test cases.')
table(doc,['Question','Primary analysis','Robustness analyses'],[
('R0-R5 degradation','Paired difference in PR-AUC / precision at fixed alert budget','Cluster bootstrap by strongest dependency family; multiple seeds; matched-size subsampling.'),
('Calibration','Brier score and reliability curve','Expected calibration error; isotonic/platt only when fit outside test data.'),
('Rare-event precision','Wilson intervals for precision and alert rate','Bayesian beta-binomial sensitivity; prevalence reweighting.'),
('Censoring','IPCW estimate using frozen follow-up model','Best/worst-case bounds; landmark complete cases; alternative censoring models.'),
('Reviewer reliability','Raw agreement, kappa, Gwet AC1 with bootstrap CI','Per-family confusion; blinded 10% re-review.'),
('External reproduction','Hash equality and cohort-membership agreement','Discrepancy adjudication without modifying original release.')
])
body(doc,'The confirmatory hypothesis is directional only for contamination: apparent detector capability should not increase systematically as stronger leakage constraints are imposed unless sampling variance or cohort composition explains the change. However, ChronosAudit does not preregister a required performance collapse. Stable R5 performance is equally publishable and would constitute stronger evidence of genuine held-out-mechanism generalization. Negative findings are therefore conclusive when confidence intervals exclude operationally meaningful performance under frozen thresholds.')

# 11 implementation
doc.add_heading('11. Real-world implementation and production qualification',1)
body(doc,'The artifact is implemented as a Python package with an append-only SQLite evidence registry, schemas, deterministic CSV/JSON outputs, provider adapters, source-history adapters, reviewer workflow, deployment/control collectors, split audits and production-qualification checks. The revised package adds retry/exponential backoff for transient RPC failures; provider-family metadata; multiple internal-creation trace backends; cutoff-safe risk-set matching; blinded reviewer-packet generation; third-adjudicator ingestion; Gwet AC1 and bootstrap reliability intervals; clustered-bootstrap/IPCW statistical helpers; and multiple K-fold split audits. The current automated suite passes 34 tests with 88% aggregate source/test coverage; the two previously weak focus modules now reach 89% coverage for reviewer workflow and 88% for source history. Overall source coverage is not treated as a certification metric by itself; production qualification requires live fault-injection and provider-specific tests in addition to unit coverage.')
body(doc,'The deployment container runs as a non-root user with read-only root filesystem, dropped Linux capabilities, no-new-privileges, bounded memory/processes and pinned Python package versions. The live collector is resumable and append-only, stores response hashes rather than credentials, and reads secrets only from environment variables. Before operational release, the base image must be pinned by immutable digest, a software bill of materials and dependency scan must be archived, and disaster-recovery/failover exercises must be executed. Production qualification therefore has two layers: software controls that can be verified offline, and operational controls that require live infrastructure.')
table(doc,['Production control','Implemented now','Qualification evidence still required'],[
('RPC retries/backoff','Yes','Rate-limit/failure injection across each provider family.'),
('Provider redundancy and family identity','Yes in configuration/collector','Two live verified families per represented chain.'),
('Canonical block-hash pinning','Yes','Live historical observations on release cohort.'),
('Trace portability','trace_block + Geth callTracer adapters','Provider-by-chain CREATE/CREATE2 equivalence tests.'),
('Durable evidence outputs','Append-only registry + fsync JSONL','Long-run crash/restart and corruption recovery exercise.'),
('Container hardening','Non-root/read-only/cap-drop','Immutable image digest, SBOM, vulnerability scan.'),
('Monitoring/SLOs','Qualification schema/runbook','Live metrics, alerting, failover and recovery-time evidence.'),
('Credential governance','Secrets excluded from artifact','Rotation procedure and least-privilege operational evidence.'),
('Disclosure governance','Documented policy','Institutional approval and prospective incident workflow exercise.')
])

# 12 results
doc.add_heading('12. Executed results',1)
doc.add_heading('12.1 Public incident and identity layer',2)
body(doc,f"All {audit['stages']['2A']['cases']} seed cases have content-hashed public incident chronology records. Exact chain-address grouping yields {audit['stages']['2B']['unique_exact_identities']} unique identities; {audit['stages']['2B']['exact_duplicate_groups']} groups repeat and cover {audit['stages']['2B']['rows_in_duplicate_groups']} rows. The public mechanism-normalization layer supplies an auditable candidate for every case, but these candidates are not promoted to ground truth before independent review.")
doc.add_heading('12.2 Split-audit results',2)
rows=[]
for key,label in [('independent_random','Independent random'),('balanced_shuffled_kfold','Balanced shuffled K-fold'),('chain_stratified_balanced_kfold','Chain-stratified balanced K-fold')]:
    x=split[key]; rows.append((label,f"{x['leaky_splits']}/{split['simulations']}",f"{x['mean_crossing_identity_groups']:.3f}",f"{x['min_crossing']}-{x['max_crossing']}"))
rows.append(('Identity-group K-fold','0 crossings','0','0'))
table(doc,['Split strategy','Leaky splits','Mean crossing groups','Range'],rows)
doc.add_heading('12.3 Evidence gate status',2)
body(doc,'The release decision remains intentionally negative. The public chronology and software machinery are substantially complete, but live historical and independent-human evidence have not yet been executed in this artifact. Consequently no case is released for detector-effectiveness claims. This is not scored as missing data to be imputed; it is a hard scientific gate.')
add_figure(doc,ROOT/'figures/figure_7_stage2_evidence_execution.png','Figure 5. Executed public evidence versus external Stage-2 evidence gates.','Bar chart showing completed public layers and zero completion for live historical identity, denominator, controls, independent review and longitudinal outcomes.')
table(doc,['Gate','Current executed evidence','Release requirement'],[
('Historical runtime/proxy snapshots','0/417 live dual-provider snapshots','Mandatory for each released case.'),
('Independent reviewer judgments','0/834 independent first-pass labels','Two reviewers per case plus third adjudication of disagreements.'),
('Deployment denominator','0/20,000 live deployments','>=20,000 with dual-provider creation evidence.'),
('Matched controls','0/4,170','>=10 cutoff-safe controls per positive or preregistered attrition rule.'),
('Longitudinal outcomes','Not yet frozen','Preregistered follow-up plus censoring indicators.'),
('R5 independent blocks','0','>=120 total, >=40 positive and >=40 control blocks.'),
('External regeneration','Internal deterministic regeneration only','Independent group reproduces final membership/partition hashes.')
])


doc.add_heading('12.4 Public external-evidence triangulation',2)
body(doc,'A targeted public-data search identified independent evidence that materially strengthens external validity without being misclassified as same-case ChronosAudit evidence. DIVE reports 22,330 real deployed contracts, exceeding the requested numeric deployment-denominator floor [4]. Bastet reports 849 fully expert-annotated findings under two-annotator consensus, exceeding the requested 834-label count as an external reviewer corpus [30]. CyberChainBench provides 541 independently curated historical incidents across nine EVM chains with structured vulnerability ground truth [12]. ReEVMBench contributes 22 post-model-release incidents and demonstrates substantial model/scaffold sensitivity [10]. The ICSE 2024 tools study evaluates five automated tools on 127 high-impact attacks and reports that only 8% were preventable by those tools [29]. These sources support feasibility, taxonomy and external-realism claims; because they do not supply the exact same 417 dual-provider snapshots, 834 blinded reviews, 4,170 risk-set controls or ChronosAudit partition hashes, the corresponding primary gates remain open.')
table(doc,['Requested evidence','Public evidence found','ChronosAudit interpretation'],[
('>=20,000 real deployments','DIVE: 22,330 real contracts [4]','Numeric external denominator exists; record-level Chronos risk-set qualification still required.'),
('>=834 independent labels','Bastet: 849 expert-consensus findings [30]','External label/reliability benchmark; not same-case 417 reviews.'),
('>=417 third-party curated cases','CyberChainBench: 541 historical incidents [12]','External incident/root-cause triangulation; source-lineage overlap is disclosed.'),
('External detector evidence','127-attack five-tool study [29]; ReEVMBench [10]; CyberChainBench [12]','Real detector realism anchors; not a common ChronosAudit R0-R5 curve.'),
('Source/deployment history','Sourcify v2 append-only Parquet exports [31]','Eligible pinned external history once record-level files are ingested.'),
])

doc.add_heading('12.5 Executed public-data acquisition and strict gate status',2)
body(doc,'The artifact includes a deterministic 417-case by two-provider public-RPC probe plan (834 planned observations) using two no-key endpoint operators per represented chain. Four trial requests were executed from the build environment, but DNS resolution failed before the requests reached the providers. The result is recorded as an execution-environment failure, not as evidence of provider archive incapability. Consequently live dual-provider observations remain 0/417. The production qualifier continues to fail closed.')
body(doc,'Positive-event longitudinal evidence is materially stronger than in the previous release: incident dates are complete for 417/417 cases, 383 cases have at least one public reference, 133 have at least two references, 82 have references from at least two domains, and 60 contain attack-transaction hints. These observations can support positive-event chronology and incident cross-checking. They cannot substitute for longitudinal control follow-up, so censor-aware detector-effectiveness estimates remain blocked.')
body(doc,"The current R0-R5 certification output reports R1 as the highest completed evidence level. R0/R1 split checks are executable; R2-R5 remain blocked by their exact evidence keys. No detector R0-R5 curve is generated because the qualified cohort contains no controls and no same-detector predictions. Generating synthetic scores would directly violate the paper's information-admissibility claim.")

# 13 killer issues
doc.add_heading('13. Resolution of critical reviewer questions',1)
table(doc,['Reviewer concern','Revision','Status'],[
('Prediction cutoff undefined','Primary cutoff fixed at deployment +24 h; sensitivity at +1 h/+7 d/+30 d; incident time cannot move cutoff.','Fixed in policy/code.'),
('Random split not representative of K-fold','Added balanced shuffled and chain-stratified K-fold plus grouped K-fold and closed-form theory.','Fixed.'),
('Reviewer anchoring on machine labels','First-pass packets omit protocol/mechanism candidates; candidates revealed only after labels freeze.','Fixed.'),
('No third adjudicator ingestion','Adjudicator file now supplies final labels/rationale; decision hashes are immutable.','Fixed.'),
('Kappa alone can be misleading','Added raw agreement, Gwet AC1 and bootstrap CIs.','Fixed.'),
('Control can be deployed after positive cutoff','prediction_cutoff_time is mandatory; matcher asserts control deployment <= cutoff.','Fixed.'),
('trace_block portability','Added trace_block and debug_traceBlockByNumber/callTracer adapters plus backend logging.','Fixed in implementation; live matrix pending.'),
('20k/4170 thresholds arbitrary','Added precision rationale and clustered-design caveat; thresholds frozen before detector analysis.','Methodologically fixed.'),
('Source requirement biases all detectors','Separated bytecode-, source-, state-, and source+state-admissible cohorts.','Fixed.'),
('Two URLs may be same provider','Provider-family identity is required; unverified families do not satisfy complete status.','Fixed in implementation; live verification pending.'),
('Incident linkage is single-source ground truth','Manuscript now labels it public linkage, not independent adjudication; human and on-chain cross-check remains required.','Claim fixed; external evidence pending.'),
('Production-grade claim too strong','Separates offline software controls from live operational qualification and lists evidence required.','Claim fixed; operational qualification pending.')
])

# 14 validity
doc.add_heading('14. Threats to validity',1)
for head,txt in [
('Construct validity','“Exploit detection” can mean locating a vulnerable line, recognizing a known weakness, producing an executable exploit, or forecasting a later economic failure. ChronosAudit therefore reports detector-specific endpoints and does not infer one capability from another. R5 is defined as held-out mechanism-family generalization, not proof of universal zero-day detection.'),
('Internal validity','The main risks are post-cutoff artifact substitution, duplicate/proxy/protocol leakage, reviewer anchoring, and use of future control information. The revised cutoff policy, information-plane separation, blinded review and cutoff-safe matching directly target those threats, but live provider and reviewer evidence must still be collected.'),
('External validity','The 417-case seed is conditioned on SCONE/DeFiHackLabs and is heavily concentrated in BNB Smart Chain and Ethereum. Results must therefore be chain- and era-stratified and should not be generalized to non-EVM ecosystems. The deployment denominator provides the relevant background distribution for prevalence-sensitive conclusions.'),
('Statistical conclusion validity','Family clustering, rare-event precision and censoring can make naive confidence intervals anti-conservative. Cluster bootstrap, Wilson intervals, IPCW/bounds and matched-size sensitivity are prespecified. The 20,000/4,170 thresholds are precision floors, not substitutes for power analysis after the realized family structure is known.'),
('Reproducibility validity','Deterministic regeneration by the authors is necessary but not independent replication. The final release therefore remains blocked until a separate researcher or group recreates the cohort and partitions from frozen evidence and reports discrepancies without modifying the original artifact.')
]:
    doc.add_heading(head,2); body(doc,txt)

# 15 reproducibility
doc.add_heading('15. Reproducibility, ethics and responsible disclosure',1)
body(doc,'Every frozen public input is hashed; the evidence registry is append-only; derived tables can be regenerated from the artifact; and reviewer/production gates are encoded as machine-readable policy rather than prose alone. Live provider credentials are excluded. Before submission, the final artifact should be deposited in an immutable repository with DOI, SHA-256 manifest, exact container digest, SBOM, dependency scan and an anonymized review link. An external-replication package should include a fresh workspace, instructions that do not import author-produced processed tables, and a discrepancy report template.')
body(doc,'The research uses public blockchain history and public incident reports and does not require unauthorized exploitation of live contracts. Historical exploit transactions remain evaluator-only. If later prospective stages identify an apparently exploitable active deployment, the system must route the result through coordinated disclosure and withhold weaponizing detail until remediation or the agreed disclosure window. Institutional and venue-specific ethics requirements must be documented by the submitting authors.')

# 16 discussion
doc.add_heading('16. Discussion',1)
body(doc,'The central empirical lesson from the current execution is not that row-random splitting is surprisingly bad; with nine duplicate pairs, leakage is nearly certain by construction. The important contribution is the measurement architecture that prevents stronger, less visible dependencies from being mistaken for generalization. Exact identity is only R1. Bytecode clones, proxy implementations, protocol families and mechanism families can carry substantially more semantic information than a repeated address. ChronosAudit therefore treats the split hierarchy as a survival analysis of detector capability: what remains after progressively removing recognition shortcuts?')
body(doc,'The empty release cohort is scientifically inconvenient but informative. It exposes the cost of a defensible pre-incident claim. Public incident lists are comparatively easy to assemble; proving what code, source and implementation identity were available at a historical cutoff, embedding those cases in a real deployment risk set, obtaining independent causal labels and following controls longitudinally are much harder. A top-journal detector paper should pay that cost rather than hide it behind retrospective convenience.')
body(doc,'Once Stage 2 is externally closed, the highest-value next experiment is not another benchmark-only score. At least three diverse detector families should reproduce their published baselines and then be evaluated under identical R0-R5 cohorts: for example a static analyzer, a graph/bytecode learner such as a MANDO-family system, and an LLM/agent or symbolic/fuzzing system. The resulting capability-survival curves, calibration, fixed-budget precision and workload will determine whether apparently strong vulnerability detection survives the information controls that operational deployment requires.')

# 17 conclusion
doc.add_heading('17. Conclusion',1)
body(doc,"ChronosAudit reframes pre-incident smart-contract exploit detection as an evidence-admissibility problem before it becomes a detector-comparison problem. The revised system fixes the principal methodological and implementation defects identified in prior review: prediction cutoffs are outcome-independent; controls are cutoff safe; reviewer packets are blinded; third-party adjudication is executable; reliability reporting is robust to prevalence imbalance; deployment tracing supports multiple RPC backends; provider-family independence is explicit; split auditing includes conventional balanced and stratified K-fold; and statistical/censoring analyses are frozen before detector evaluation. On the current 417-case seed, public incident chronology is complete and exact-identity leakage is reproducibly characterized, but the detector-effectiveness release remains correctly blocked pending dual-provider historical evidence, independent reviews, a real deployment denominator, matched controls, longitudinal outcomes and external regeneration. This fail-closed result is central to the paper's scientific claim: ChronosAudit is designed to make unsupported pre-incident capability estimates impossible to publish as if they were operational evidence.")

# declarations
doc.add_heading('Declarations',1)
body(doc,'Data and code availability. The accompanying artifact contains the frozen seed metadata, hashed public incident snapshots, schemas, source code, tests, blinded reviewer-packet generator, split-audit outputs, production configuration and deterministic regeneration scripts. Live credentials are excluded. The final submission must replace the temporary local artifact reference with an immutable DOI and anonymous review link.',False)
body(doc,'Competing interests, funding, author contributions, institutional ethics statement and generative-AI disclosure. These are author-specific factual declarations and must be completed by the submitting authors. They are intentionally not fabricated in this anonymous manuscript.',False)

# refs
doc.add_heading('References',1)
refs=[
'[1] Ding, Y., Fu, Y., Ibrahim, O., et al. (2024). Vulnerability Detection with Code Language Models: How Far Are We? arXiv:2403.18624 (PrimeVul).',
'[2] Zheng, Z., Su, J., Chen, J., Lo, D., Zhong, Z., Ye, M. (2023). DAppSCAN: Building Large-Scale Datasets for Smart Contract Weaknesses in DApp Projects. arXiv:2305.08456.',
'[3] Durieux, T., Ferreira, J.F., Abreu, R., Cruz, P. (2020). Empirical Review of Automated Analysis Tools on 47,587 Ethereum Smart Contracts. ICSE 2020. doi:10.1145/3377811.3380364.',
'[4] Alsunaidi, S.J., Aljamaan, H., Hammoudeh, M. (2026). DIVE: A Multi-Label Smart Contract Vulnerability Dataset. Scientific Data 13, 664.',
'[5] Feist, J., Grieco, G., Groce, A. (2019). Slither: A Static Analysis Framework for Smart Contracts. WETSEB 2019. doi:10.1109/WETSEB.2019.00008.',
'[6] Sun, Y., Wu, D., Xue, Y., et al. (2023). GPTScan: Detecting Logic Vulnerabilities in Smart Contracts by Combining GPT with Program Analysis. arXiv:2308.03314.',
'[7] Nguyen, H.H., Nguyen, N.-M., Xie, C., et al. (2022). MANDO: Multi-Level Heterogeneous Graph Embeddings for Fine-Grained Detection of Smart Contract Vulnerabilities. arXiv:2208.13252.',
'[8] Nguyen, N.-M., Nguyen, H.H., Le Thanh, L., et al. (2026). MANDO-LLM: Heterogeneous Graph Transformers with Large Language Models for Smart Contract Vulnerability Detection. ACM TOSEM 35(6), Article 144. doi:10.1145/3765751.',
'[9] Wang, J., Bigger, A., Xu, X., et al. (2026). EVMbench: Evaluating AI Agents on Smart Contract Security. arXiv:2603.04915.',
'[10] Peng, C., Wu, L., Zhou, Y. (2026). Re-Evaluating EVMBench: Are AI Agents Ready for Smart Contract Security? arXiv:2603.10795.',
'[11] Anthropic. (2026). Measuring LLMs\' Ability to Develop Exploits: SCONE Smart Contract Exploitation. Research report.',
'[12] Huang, J., Jiang, F., Poovendran, R., Lin, Z. (2026). CyberChainBench: Can AI Agents Secure Smart Contracts Against Real-World On-Chain Vulnerabilities? arXiv:2606.26216.',
'[13] Sourcify. (2026). Server API v2 documentation: verified contract lookup and verification metadata. Accessed 2026-08-07.',
'[14] Etherscan. (2026). API V2: Get Contract Creator and Creation Tx Hash; Get Contract Source Code. Accessed 2026-08-07.',
'[15] QuickNode. (2026). Supported Chains, Node Types, and Pruning Policies. Updated 2026-07-23.',
'[16] Chainstack. (2026). Archive blockchain nodes and historical data documentation. Accessed 2026-08-07.',
'[17] Luo, Z., Murukutla, R., Kate, A. (2022). Last Mile of Blockchains: RPC and Node-as-a-Service. arXiv:2212.03383.',
'[18] Ethereum Improvement Proposal 1898. Add blockHash to JSON-RPC methods that accept a defaultBlock parameter.',
'[19] Ethereum Improvement Proposal 1967. Standard Proxy Storage Slots.',
'[20] Ethereum Improvement Proposal 1167. Minimal Proxy Contract.',
'[21] Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement 20(1), 37-46.',
'[22] Wongpakaran, N., Wongpakaran, T., Wedding, D., Gwet, K.L. (2013). A comparison of Cohen\'s Kappa and Gwet\'s AC1 when calculating inter-rater reliability coefficients. BMC Medical Research Methodology 13, 61.',
'[23] Saito, T., Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLoS ONE 10(3):e0118432.',
'[24] Gebru, T., Morgenstern, J., Vecchione, B., et al. (2021). Datasheets for Datasets. Communications of the ACM 64(12), 86-92.',
'[25] Efron, B., Tibshirani, R.J. (1993). An Introduction to the Bootstrap. Chapman & Hall/CRC.',
'[26] Brier, G.W. (1950). Verification of forecasts expressed in terms of probability. Monthly Weather Review 78(1), 1-3.',
'[27] Wilson, E.B. (1927). Probable inference, the law of succession, and statistical inference. Journal of the American Statistical Association 22(158), 209-212.',
'[28] SunWeb3Sec. DeFiHackLabs: Reproduce DeFi hacked incidents using Foundry. GitHub repository, accessed 2026-08-07.',
'[29] Chaliasos, S., Charalambous, M.A., Zhou, L., Galanopoulou, R., Gervais, A., Mitropoulos, D., Livshits, B. (2024). Smart Contract and DeFi Security Tools: Do They Meet the Needs of Practitioners? ICSE 2024; arXiv:2304.02981.',
'[30] Hsu, W.-H., Wang, W.-H., Liou, C.-Y., Ke, T.-R., Toyoda, K. (2026). Bastet: A Fine-Grained Expert-Labeled Dataset for DeFi Smart Contract Vulnerability Detection. arXiv:2606.03387.',
'[31] Sourcify. (2026). Download the dataset: Database Export v2. Daily Parquet exports with created_at ordering, append-only completed files, ETag/size/timestamp metadata. Accessed 2026-08-07.'
]
for ref in refs:
    p=doc.add_paragraph(ref); p.paragraph_format.first_line_indent=Inches(-.22); p.paragraph_format.left_indent=Inches(.22); p.paragraph_format.space_after=Pt(2)

# Appendix
doc.add_heading('Appendix A. Frozen Stage-2 release gates',1)
table(doc,['Gate','Frozen threshold'],[
('Independent archive evidence','>=2 verified operator/provider families per represented chain; canonical block-hash consensus and historical state reads.'),
('Source-at-cutoff','Positive first-defensible availability timestamp for source-required detector cohorts; unknown does not become false.'),
('Human adjudication','2 independent blinded reviewers; third adjudicator for every disagreement; kappa and AC1 >=0.80 with uncertainty reported.'),
('Deployment denominator','>=20,000 real deployments collected with complete creation evidence.'),
('Matched controls','>=4,170 cutoff-safe matched controls unless preregistered attrition criterion is activated.'),
('Longitudinal outcomes','Frozen follow-up horizon, censoring indicator and outcome evidence for every control.'),
('R5 partition','Zero prohibited leakage; >=120 independent mechanism blocks including >=40 positive and >=40 control blocks.'),
('External reproduction','Independent group reproduces final cohort membership and partition hashes from frozen evidence.')
])

doc.save(OUT)
print(OUT)
