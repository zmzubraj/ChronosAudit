from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Inches
from pathlib import Path
p=Path('/mnt/data/chronos_stage2_external/paper/ChronosAudit_PreIncident_TopJournal_Stage2_ExternalEvidence_Revised.docx')
doc=Document(p)

def after(par, text):
    new_p = OxmlElement('w:p')
    par._p.addnext(new_p)
    q=Paragraph(new_p, par._parent)
    q.style='Body Text'
    q.paragraph_format.first_line_indent=Inches(.22)
    q.add_run(text)
    return q

def insert_after_contains(needle,text):
    for par in doc.paragraphs:
        if needle in par.text:
            after(par,text); return True
    print('NOT FOUND',needle); return False

insert_after_contains('Source admissibility is asymmetric.',
    'For reproducible source-history reconstruction, the implementation also ingests pinned Sourcify Database Export v2 Parquet/CSV records. Sourcify documents daily table exports ordered by created_at, an append-only pattern for completed files, and storage metadata including ETag, object size and modification time [31]. ChronosAudit stores the upstream object identity/checksum alongside normalized observations, so a source-at-cutoff conclusion can be regenerated from a frozen export rather than relying only on today\'s API response. A verification timestamp at or before the cutoff is positive evidence; later verification or absence from Sourcify never proves universal historical non-availability.')
insert_after_contains('Agreement reporting is deliberately redundant.',
    'The executable review workflow additionally reports nominal Krippendorff alpha and supports an external-corroboration table after the two first-pass labels are frozen. Bastet supplies 849 fully expert-annotated DeFi findings under a two-annotator consensus workflow [30], while CyberChainBench supplies structured ground truth for 541 historical incidents [12]. These corpora can stress-test taxonomy mappings and disagreement patterns, but they are never counted as the two blinded same-case ChronosAudit reviewers. The distinction between independent curator, independent reviewer and independent incident-source lineage is stored explicitly.')
insert_after_contains('The primary denominator target is at least 20,000 deployments',
    'The public-data feasibility check is positive at the numeric level: DIVE reports 22,330 real Ethereum smart contracts deployed during 2016-2024 [4], exceeding the 20,000-contract floor. This external corpus is treated as a denominator/control candidate frame only. It becomes a ChronosAudit risk set only after record-level address, deployment time, cutoff-admissible evidence and follow-up are materialized; this avoids turning a large but differently sampled dataset into false controls.')
insert_after_contains('For censoring-sensitive analyses, inverse-probability-of-censoring weighting',
    'Censoring is now executable in the artifact rather than specified only in prose. The statistics module implements Kaplan-Meier estimation of the censoring survival function, IPCW binary metrics, landmark outcome extraction, and best/worst-case partial-identification bounds. The positive-event side is materially populated: 417/417 incident dates are frozen, 383 cases have at least one public reference, 133 have at least two references, 82 have references from at least two domains, and 60 expose attack-transaction hints. These facts close positive chronology but do not create control non-events; controls remain censored until an actual follow-up horizon is observed.')
insert_after_contains('The current seed contains 9 repeated exact-identity groups',
    'A strict machine-readable R0-R5 evidence certificate accompanies the release. On the current 417-row Stage-2 table, R0 and R1 are certified; R2 is blocked by incomplete outcome-independent prediction-cutoff observations, R3 by missing final protocol-family adjudication, R4 by incomplete implementation-family reconstruction, and R5 by missing final mechanism-family adjudication. Preliminary mechanism/protocol candidates are deliberately ignored by the certifier. The highest certified level is therefore R1, and no R0-R5 detector capability-survival curve is claimed until a common positive/control cohort and detector predictions exist.')
# Add another sentence to implementation with coverage if absent
insert_after_contains('The current automated suite passes 34 tests',
    'Focused coverage is 89% for review_workflow.py and 88% for source_history.py; deployment_stream.py is 93% and split_audit.py is 100%. Passing tests are not treated as live production evidence: the production qualifier separately requires provider-family observations, historical snapshots, source/deployment records, completed adjudication, denominator, controls, censor-aware outcomes and R5 blocks.')
doc.save(p)
print(p)
